from __future__ import annotations

from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path("/Users/j1ng/Tools/sunny")
OUT = ROOT / "outputs" / "思远物流系统测试版用户手册.docx"

BLUE = "2457C5"
DARK = "12233A"
MUTED = "667085"
LIGHT_BLUE = "EAF1FF"
LIGHT_GRAY = "F6F8FB"
BORDER = "D8E0EA"
GREEN = "EAF8EF"
WARN = "FFF7E6"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color: str = BORDER, size: str = "8") -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = tc_pr.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tc_pr.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_table_borders(table, color: str = BORDER) -> None:
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(cell, color)


def set_cell_text(cell, text: str, bold: bool = False, color: str = DARK, size: int = 9) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table)
    if widths:
        for idx, width in enumerate(widths):
            for cell in table.columns[idx].cells:
                cell.width = Cm(width)
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, LIGHT_BLUE)
        set_cell_text(cell, header, bold=True, color=DARK, size=9)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], value, color=DARK, size=8)
    doc.add_paragraph()


def add_callout(doc: Document, title: str, text: str, fill: str = LIGHT_BLUE) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Cm(16.2)
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_border(cell, "A9C3FF")
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(title)
    r.bold = True
    r.font.name = "Arial"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor.from_string(BLUE)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(2)
    r2 = p2.add_run(text)
    r2.font.name = "Arial"
    r2._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    r2.font.size = Pt(9)
    r2.font.color.rgb = RGBColor.from_string(DARK)
    doc.add_paragraph()


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        p.add_run(item)


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(3)
        p.add_run(item)


def set_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(DARK)
    normal.paragraph_format.line_spacing = 1.18
    normal.paragraph_format.space_after = Pt(6)

    for style_name, size, color, before, after in [
        ("Heading 1", 18, BLUE, 14, 8),
        ("Heading 2", 14, DARK, 10, 5),
        ("Heading 3", 11, DARK, 8, 3),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name in ("List Bullet", "List Number"):
        style = doc.styles[style_name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(10)
        style.paragraph_format.left_indent = Cm(0.55)
        style.paragraph_format.first_line_indent = Cm(-0.25)


def add_title(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("思远物流系统测试版用户手册")
    r.font.name = "Arial"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    r.font.size = Pt(24)
    r.font.bold = True
    r.font.color.rgb = RGBColor.from_string(DARK)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run("面向测试用户的工作流、模块功能与反馈指南")
    r.font.name = "Arial"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor.from_string(MUTED)

    add_callout(
        doc,
        "测试定位",
        "当前系统处于可用雏形阶段，核心目标是验证真实业务流程、角色权限、数据落地和页面易用性。请测试人员重点反馈：流程是否符合实际工作、字段是否缺失、按钮是否出现在正确时机、数据刷新后是否保留。",
    )

    add_table(
        doc,
        ["文档项", "说明"],
        [
            ["系统名称", "思远物流 AI TMS / OMS"],
            ["文档版本", "测试版 v0.1"],
            ["生成时间", datetime.now().strftime("%Y-%m-%d %H:%M")],
            ["适用对象", "管理员、业务员、客服、仓库、财务、测试用户"],
            ["测试地址", "以当前部署地址为准；如地址变化，由管理员另行通知"],
        ],
        [4.0, 12.0],
    )


def build_doc() -> None:
    doc = Document()
    set_styles(doc)
    add_title(doc)

    doc.add_heading("1. 系统总体说明", level=1)
    add_callout(
        doc,
        "一句话理解系统",
        "系统围绕跨境物流订单，从仓库入库、业务员建单、客服审核、渠道排货、仓库打单出库、客服跟进轨迹、财务收付款，到报价查价和基础资料维护，形成一套可持续扩展的 TMS / OMS 工作台。",
        GREEN,
    )
    add_bullets(
        doc,
        [
            "当前版本已经不是纯前端 Demo，关键数据逐步改为后端 API 和数据库持久化，刷新页面后应保留。",
            "报价查价、仓库包裹、运单、权限、登录日志、操作日志等核心数据已开始后端落地。",
            "外部 API 暂以预留或模拟为主，例如仓库系统包裹 API、承运商转单 API、轨迹抓取 API，后续可继续接入真实服务。",
            "AI 能力以辅助建议为主，用于降低录单、排货、跟进和异常处理时的信息整理成本。",
        ]
    )

    doc.add_heading("2. 用户角色与数据权限", level=1)
    add_table(
        doc,
        ["角色", "主要职责", "当前可测能力", "数据范围"],
        [
            ["管理员", "系统配置、权限、基础资料、报价管理、全局数据查看", "全部模块、权限配置、报价表管理、加价规则、系统设置", "全局数据"],
            ["业务员", "创建出货订单、查询报价、跟进自己名下客户", "订单履约、渠道排货、报价查询、基础资料查看", "只能看自己名下客户的数据"],
            ["客服", "订单审核、轨迹跟进、问题件处理、客户沟通", "订单审核、修改、轨迹、问题件、报价查询", "按客服业务范围查看"],
            ["仓库", "入库、包裹明细、合票、打单、出库", "仓库管理、面单队列、出库动作、仓库数据查看", "可看所有业务员名下客户的仓库相关数据"],
            ["财务", "收款、应收应付、对账、流水", "财务结算、费用查看、收付款、核销类动作", "财务相关数据"],
            ["客户", "客户门户、本人订单与费用", "本人运单、本人问题件、本人费用、价格查询", "本人客户数据"],
        ],
        [2.2, 4.0, 5.2, 4.4],
    )
    add_callout(
        doc,
        "权限测试重点",
        "测试时请重点关注：业务员是否看不到其他业务员客户；仓库是否可以看到所有业务员的包裹；业务员是否看不到报价成本、代理加价、价格表管理；财务是否可以进入财务模块但不能修改系统权限。",
        WARN,
    )

    doc.add_heading("3. 核心订单全链路流程", level=1)
    add_numbered(
        doc,
        [
            "仓库接收包裹数据：包裹信息来自仓库系统 API 或当前测试数据，按客户单号聚合展示。",
            "业务员整理货物并创建出货订单：在订单履约中录入客户、目的地、件数、重量、备注等基础信息。",
            "订单进入待审核：客服或有权限人员检查资料是否完整，选择审核通过或审核不通过。",
            "审核通过进入待排货：渠道排货模块为订单分配代理与渠道。",
            "排货后进入待出库：仓库在面单队列处理打单和出库。",
            "仓库出库后进入已出库：客服补齐转单号，补齐后自动进入待离港。",
            "待离港阶段填写 ETD / ETA：确认离港前必须填写预计离港和预计到港信息。",
            "后续状态流转：已离港、已到港、已派送、已签收；已到港可直接进入已签收。",
            "异常分支：待离港、已到港、已派送阶段可以转问题件，需填写问题类型与说明。",
            "全程记录日志：审核、排货、出库、转单号、轨迹、收款、删除等关键动作写入操作日志。",
        ]
    )
    add_table(
        doc,
        ["状态", "业务含义", "主要操作人", "下一步"],
        [
            ["待审核", "业务员创建订单后等待资料审核", "客服 / 管理员", "审核通过进入待排货；审核不通过退回修改"],
            ["审核不通过", "资料或信息不满足出货要求", "业务员 / 客服", "修改后重新提交审核"],
            ["待排货", "需要选择代理和渠道", "业务员 / 排货人员", "分配渠道后进入待出库"],
            ["待出库", "等待仓库打单和出库", "仓库", "打单、出库后进入已出库"],
            ["已出库", "货物已交出，需要补齐快递转单号", "客服", "填写转单号后进入待离港"],
            ["待离港", "等待航班/船期/干线离港", "客服 / 业务员", "填写 ETA/ETD 后进入已离港，或转问题件"],
            ["已离港", "货物已离开发货地", "客服 / 业务员", "跟进到港"],
            ["已到港", "货物到达目的港或目的地节点", "客服 / 业务员", "进入已派送或直接签收，或转问题件"],
            ["已派送", "末端派送中", "客服 / 业务员", "确认签收，或转问题件"],
            ["已签收", "订单完成", "客服 / 业务员", "归档、财务核对"],
            ["问题件", "订单存在查验、异常、资料问题或服务问题", "客服 / 业务员", "处理后按实际情况恢复流程"],
        ],
        [2.5, 5.0, 3.2, 5.0],
    )

    doc.add_heading("4. 模块功能说明", level=1)

    modules = [
        (
            "运营工作台",
            "全系统订单状态看板，用于快速查看订单总量、状态池和当前需要处理的订单。",
            [
                "状态池：全部、待审核、审核不通过、待排货、待出库、已出库、待离港、已离港、已到港、已派送、问题件、已签收。",
                "订单列表：展示创建时间、客户名称、运单号、转单号、目的地、渠道、代理、业务员归属、收款、状态、时效等字段。",
                "列设置：可按岗位调整列顺序，例如客户优先、代理优先或自定义顺序。",
                "订单详情：点击运单号可打开基础订单信息，查看单号、客户、渠道、代理、费用、轨迹、备注等。",
            ],
            "测试建议：切换状态池、调整列设置、点击运单号查看详情，确认字段是否足够业务使用。",
        ),
        (
            "仓库管理",
            "仓库岗位的第一工作区，处理入库包裹、包裹明细、合票、面单队列和待仓库出货。",
            [
                "包裹明细：按客户单号聚合展示，支持搜索客户单号、快递单号和客户编码。",
                "字段：客户单号-快递单号、件数、实重、尺寸、方数、5000 材积、6000 材积、扫描时间、备注、异常。",
                "合票：支持仅合并包裹，以及合并并创建出货单两种情况。",
                "面单队列&待仓库出货：出货单审核通过后进入，仓库执行打单和出库。",
                "收货交接单：从渠道排货等环节抓取必要字段，用于输出交接资料，后续可支持 PDF 或 Word 下载。",
            ],
            "测试建议：用客户单号搜索包裹，查看部分到仓件数；测试备注填写；测试多个包裹合并后能查看原始明细。",
        ),
        (
            "订单履约",
            "业务员录单和客服审核入口，主要处理订单创建、审核、修改、收款、日志和删除。",
            [
                "新建出货订单：录入客户、目的地、件数、重量、渠道、备注等基础信息。",
                "审核：支持审核通过和审核不通过，审核通过后进入待排货。",
                "修改：统一处理状态、转单号、轨迹、备注等人工修正内容。",
                "收款：记录收款金额、币种和收款方式。",
                "操作日志：查看该订单全生命周期关键动作。",
                "删除：重要操作需要二次确认，删除后刷新不应再出现。",
            ],
            "测试建议：创建订单后刷新页面，确认订单仍存在；测试审核通过后是否进入渠道排货；测试收款币种和方式是否保存。",
        ),
        (
            "渠道排货",
            "排货人员或业务员为待排货订单指定代理和渠道。",
            [
                "只处理待排货订单，避免非必要按钮出现在错误状态。",
                "分配渠道：可从基础资料选择代理和渠道，也可手动输入。",
                "排货日志：记录分配渠道、修改渠道、操作人员和时间。",
                "分配完成后状态进入待出库，进入仓库管理的面单队列/待仓库出货流程。",
            ],
            "测试建议：选择待排货订单分配代理渠道，刷新确认状态不丢；查看排货日志是否记录完整。",
        ),
        (
            "轨迹监控 / 客服跟进",
            "用于客服和业务员跟进已出库后的物流动态。",
            [
                "已出库：客服补齐转单号，补齐后进入待离港。",
                "待离港：填写 ETD 和 ETA 后才能进入已离港。",
                "已到港和已派送：可继续推进状态，也可转问题件。",
                "轨迹导入：支持通过表格导入轨迹，系统按单号取最新日期的轨迹描述。",
            ],
            "测试建议：测试无 ETA/ETD 是否不能确认离港；测试填写转单号后状态是否自动进入待离港。",
        ),
        (
            "问题件中心",
            "集中处理订单异常、客户问题和内部服务链路异常。",
            [
                "新建问题：关联运单，填写问题类型和说明。",
                "回复/查看：记录内部或客户可见回复。",
                "关闭问题：问题处理完成后关闭，关闭后原则上不继续回复。",
                "权限：客户只看本人问题件，员工按角色查看处理。",
            ],
            "测试建议：从订单状态转问题件，确认问题说明保留；测试关闭后是否禁止继续处理。",
        ),
        (
            "报价查价",
            "用于业务员查询报价，管理员维护价格表和加价规则。",
            [
                "业务员视角：只看到最终报价、单价、渠道、时效和 Top 推荐，不显示成本价、代理加价、毛利。",
                "管理员视角：可导入价格表、维护备注、管理代理/渠道/国家加价规则。",
                "价格表管理：上传 XLS 后落地到后端，刷新不丢；备注由管理员手工维护。",
                "渠道加价：可按代理统一加价，也可点击线路详情对某条真实渠道/小表进行自定义加价。",
                "查价结果：可输出最便宜 Top3 和最快 Top3；渠道结果来自价格表匹配。",
            ],
            "测试建议：分别用管理员和业务员查看报价页，确认敏感字段是否隐藏；导入价格表后刷新确认仍在。",
        ),
        (
            "财务结算",
            "处理应收、应付、收款、付款、对账、核销和流水。",
            [
                "应收费用：业务实际收客户的钱，业务和财务都应能核对。",
                "应付费用：实际付给代理的费用，代理名称默认来自订单代理，可修改；业务员一般不进入该大模块。",
                "业务成本：含公司运营成本在内的内部业务费用，可用于利润核算。",
                "对账单：客户对账、代理对账和收付款记录。",
                "操作要求：费用调整、核销、收付款等高风险动作需要二次确认并记录日志。",
            ],
            "测试建议：对一个订单生成费用，录入收款，查看财务详情中的应收、应付、成本和利润是否合理。",
        ),
        (
            "基础资料",
            "维护客户、代理、渠道、国家、费用名称、汇率等基础数据。",
            [
                "客户资料：客户编码、客户简称、客户全称、客户类型、业务员归属。",
                "代理资料：代理编码、代理简称、代理名称。",
                "渠道资料：后续可补充材积除数、进位规则、渠道类型等配置。",
                "原则：默认软删除或停用，不做危险物理删除。",
            ],
            "测试建议：新增客户时填写业务员归属，确认业务员账号只看到自己名下客户订单。",
        ),
        (
            "系统设置",
            "管理员维护员工账号、角色权限、AI 接口安全、审计和系统基础配置。",
            [
                "角色权限分配：按管理员、客服、业务员、仓库、财务、客户进行授权。",
                "权限边界：报价模块特殊，业务员仅可查价，不可看价格表管理、加价和成本。",
                "员工账号管理：可新增员工、重置密码、查看角色边界。",
                "审计：高风险操作需要写入审计日志。",
            ],
            "测试建议：修改某角色权限后重新登录，确认菜单和接口权限都生效，而不只是前端隐藏。",
        ),
        (
            "个人中心与登录安全",
            "每个登录用户可查看自己的账号信息、修改密码和登录日志。",
            [
                "登录页：支持图片验证码，回车可登录。",
                "错误提示：密码错误显示用户名或密码错误；验证码错误显示验证码错误。",
                "登录日志：记录登录时间、IP、地区和设备。",
                "退出登录：位于系统外层顶部入口，不放在个人中心内部。",
            ],
            "测试建议：故意输错密码和验证码，确认提示区分清楚；查看登录日志是否新增记录。",
        ),
    ]

    for name, purpose, features, suggestion in modules:
        doc.add_heading(name, level=2)
        p = doc.add_paragraph()
        p.add_run("模块定位：").bold = True
        p.add_run(purpose)
        add_bullets(doc, features)
        add_callout(doc, "测试建议", suggestion, LIGHT_GRAY)

    doc.add_heading("5. AI 工作流说明", level=1)
    add_bullets(
        doc,
        [
            "AI 当前定位为辅助，不直接替代人工决策。",
            "可用于履约助手、报价解释、异常识别、客户沟通草稿、问题件归类等场景。",
            "AI 输出应作为建议，关键业务动作仍需要人工确认。",
            "AI 接口调用需要登录用户权限，避免匿名调用。",
        ]
    )

    doc.add_heading("6. 测试用户建议覆盖的场景", level=1)
    add_table(
        doc,
        ["测试场景", "操作步骤", "期望结果"],
        [
            ["角色隔离", "分别登录管理员、业务员、仓库、财务", "菜单、按钮、接口数据范围符合角色职责"],
            ["业务员数据范围", "给不同客户配置不同业务员归属后登录业务员账号", "只能看到自己名下客户订单"],
            ["订单创建", "业务员新建出货订单并刷新页面", "订单仍存在，状态为待审核"],
            ["审核流转", "客服审核通过和审核不通过", "通过后到待排货；不通过保留原因"],
            ["排货", "在渠道排货分配代理和渠道", "状态进入待出库，排货日志有记录"],
            ["仓库出库", "在仓库管理打单并出库", "状态进入已出库，面单信息可查看"],
            ["转单号", "已出库订单填写转单号", "自动进入待离港"],
            ["ETA/ETD", "待离港订单尝试确认离港", "未填写 ETA/ETD 时被阻止；填写后可进入已离港"],
            ["报价查价", "业务员查询报价，管理员导入价格表和维护加价", "业务员不看到成本/毛利，管理员可管理规则"],
            ["财务", "录入收款、生成费用、查看财务详情", "金额、币种、方式、利润结构清楚"],
            ["删除/高危动作", "执行删除、出库、签收、核销等动作", "出现二次确认，日志可追踪"],
        ],
        [3.0, 6.2, 6.8],
    )

    doc.add_heading("7. 当前已知边界与后续待接入", level=1)
    add_bullets(
        doc,
        [
            "仓库 XLS 后续会通过仓库系统 API 自动同步；当前版本已用真实样例数据作为测试源。",
            "承运商/代理 API、自动获取转单号、自动轨迹抓取仍属于后续集成范围。",
            "电子秤、PDA、打印客户端、微信入口、客户门户深度功能仍可继续扩展。",
            "部分模块仍处于快速迭代期，测试重点应放在流程是否符合实际业务，而非最终视觉定稿。",
            "若遇到按钮能点但后端拒绝、字段缺失、状态不符合业务，请记录账号、订单号、操作路径和截图。",
        ]
    )

    doc.add_heading("8. 反馈记录模板", level=1)
    add_table(
        doc,
        ["字段", "填写说明"],
        [
            ["测试角色", "例如：业务员、客服、仓库、财务、管理员"],
            ["所在模块", "例如：订单履约 / 仓库管理 / 报价查价"],
            ["订单号或客户单号", "尽量填写，便于定位数据"],
            ["实际操作", "按步骤描述点击了什么、输入了什么"],
            ["实际结果", "页面显示、报错、状态变化或数据结果"],
            ["期望结果", "用户认为应该如何显示或流转"],
            ["优先级", "高：阻断流程；中：影响效率；低：体验优化"],
            ["截图/录屏", "如有，请附上"],
        ],
        [4.0, 12.0],
    )

    doc.add_heading("9. 给测试用户的简短说明", level=1)
    add_callout(
        doc,
        "建议测试方式",
        "请不要只看单个页面是否好看，而是按一个真实订单从入库、建单、审核、排货、出库、补转单号、离港、到港、派送、签收完整走一遍。任何不符合实际岗位习惯的地方，都可以直接反馈。",
        GREEN,
    )

    doc.save(OUT)


if __name__ == "__main__":
    build_doc()
    print(OUT)
